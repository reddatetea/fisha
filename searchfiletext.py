# _*_ conding:utf-8 _*_

'''
搜索给定的文本内容
选取从文件夹或文件开始
如果是文件夹，则查找该路径下（包含子文件夹内）所有的指定格式（包括py,ipynb,txt,doc,docx,xls,xlsx,pdf等后缀名)文本文件
，并把所有包含文本内容的行及文件名打印出来
'''

import os
import easygui
import openpyxl
import xlrd
import docx
import pdfplumber
import csv

def openPyIpynbTxt(file,fname,fmingxi):
    try :
        with open(file, encoding='utf-8', errors='replace') as f:
            for j in f.readlines():
                if keyname in j:
                    fmingxi.write('\n')
                    fmingxi.write('%s\n' % file)
                    fmingxi.write('%s\n' % j)
                else:
                    continue
    except :
        easygui.msgbox(f'{file}有点问题，请予以关注')

    return fname,file

def openXlsXlsx(file,fname,fmingxi):
    try :
        if file.startswith('~$') == False:
            if file.lower().endswith('.xlsx'):
                wb = openpyxl.load_workbook(file, data_only=True)
                for ws in wb.worksheets:
                    for index, row in enumerate(ws.values):
                        i = row
                        for j in i:
                            if keyname in str(j):
                                fmingxi.write('\n')
                                fmingxi.write(f'{file}\n')
                                fmingxi.write(f'{ws.name}\n')
                                fmingxi.write(f'{str(i)}\n')
                            else:
                                continue
            else:
                wb = xlrd.open_workbook(file) # 打开xls文件
                for ws in wb.sheets():
                    nrows = ws.nrows # 获取表的行数
                    for i in range(nrows): 
                        for j in ws.row_values(i):
                            if keyname in str(j):
                                fmingxi.write('\n')
                                fmingxi.write(f'{file}\n')
                                fmingxi.write(f'{ws.name}\n')
                                fmingxi.write(f'{str(ws.row_values(i))}\n')
                            else:
                                continue
        else:
            pass

    except :
        easygui.msgbox(f'{file}有点问题，请予以关注')
    return fname,file

def openDocDocx(file,fname,fmingxi):
    try:
        document = docx.Document(file)
        for paragraph in document.paragraphs:
            if keyname in paragraph.text:
                fmingxi.write('\n')
                fmingxi.write('%s\n' % file)
                fmingxi.write('%s\n' % paragraph.text)
            else:
                continue
        # 在表格中搜索
        tables = document.tables
        for table in document.tables:
            rows = table.rows
            for row in rows:
                cells = row.cells
                for cell in row.cells:
                    if keyname in cell.text:
                        hang = ''.join([cell.text for cell in row.cells])
                        fmingxi.write('\n')
                        fmingxi.write('%s\n' % file)
                        fmingxi.write('%s\n' % hang)
    except:
        easygui.msgbox(f'{file}有点问题，请予以关注')
    return fname,file

def openPdf(file,fname,fmingxi):
    try:
        pdf = pdfplumber.open(file)
        for page in pdf.pages:
            content = page.extract_text()
            hangs = content.split('\r\n')
            for hang in hangs:
                if keyname in hang:
                    fmingxi.write('\n' )
                    fmingxi.write('%s\n' % file)
                    fmingxi.write('%s\n' % str(hang))
    except:
        easygui.msgbox(f'{file}有点问题，请予以关注')
    return fname,file

def openCsv(file,fname,fmingxi):
    try:
        with open(file, 'r', encoding='utf-8-sig', newline='') as f:  # utf-8不行 ,unicode_escape
            f_csv = csv.reader(f)
            for row in f_csv:
                hang = [j.strip() for j in row]
                hang_string = ''.join([j.strip() for j in row])
                if keyname in hang_string:
                    fmingxi.write('\n')
                    fmingxi.write('%s\n' % file)
                    fmingxi.write('%s\n' % str(hang))
                else:
                    continue
    except:
        easygui.msgbox(f'{file}有点问题，请予以关注')
    return fname,file

def searchSingleFile(fname,file):
    suffix = os.path.splitext(file)[1].lower()
    try:
        with open(fname, 'a', encoding='utf-8', errors='replace') as fmingxi:
            if suffix in ['.py']:
                openPyIpynbTxt(file, fname, fmingxi)
            elif suffix in ['.ipynb']:
                openPyIpynbTxt(file, fname, fmingxi)
            elif suffix in ['.txt']:
                openPyIpynbTxt(file, fname, fmingxi)
            elif suffix in ['.csv']:
                openCsv(file, fname, fmingxi)
            elif suffix in ['.pdf']:
                openPdf(file, fname, fmingxi)
            elif suffix in ['.doc', '.docx']:
                openDocDocx(file, fname, fmingxi)
            elif suffix in ['.xls', '.xlsx']:
                openXlsXlsx(file, fname, fmingxi)
            else:
                easygui.msgbox('暂不支持此后缀的搜索')
    except:
        pass

    return fname,file

def search_file(fname,start_dir, target,keyname):
    os.chdir(start_dir)
    lst = [i for i in os.listdir(os.curdir) if '~$' not in i]
    for each_file in lst:      #获取当前工作目录os.curdir下的文件列表！
        if  os.path.isfile(each_file) :
            file_list = []
            ext = os.path.splitext(each_file)[1]     #文件列表每个文件，os.path.splitext分割为路径和文件名,后缀
            if ext in target :
                file_list.append(os.getcwd() + os.sep + each_file + os.linesep)  # 使用os.sep是程序更标准
                file = os.getcwd() + os.sep + each_file
                fname,file = searchSingleFile(fname,file)
            else :
                continue
        else:
            search_file(fname,each_file, target,keyname)  # 递归调用
            try :
                os.chdir(os.pardir)  # 递归调用后切记返回上一层目录
            except:
                easygui.msgbox('已搜索到电脑的根目录')
  
msg = '请输入要查询文本的关键字'
keyname = easygui.enterbox(msg)
program_dir = os.getcwd()
searchfile_path = r'd:\searchfile'
if not os.path.exists(searchfile_path):
    os.makedirs(searchfile_path)
keyname1 = keyname
fuhao = r':\/?*[]<>"'
for j in fuhao:
    if j in keyname:
        keyname1 = keyname.replace(j, '')
    else:
        continue
fname = searchfile_path + os.sep + 'search_%s.txt' % keyname1
if os.path.exists(fname):          #删除以前的搜索文件
    os.remove(fname)
msg = '请点选查询的起点："文件夹" 或 "文件"'
choice = easygui.choicebox(msg,title = 'dir or file',choices = ['文件夹','文件'])
if choice == '文件夹':
    msg = '请点选要要查找文件所在文件夹'
    start_dir = easygui.diropenbox(msg)
    msg = '请输入文件后缀'
    target0 = easygui.multchoicebox(msg, title='文件后缀',choices=['py', 'ipynb', 'txt', 'pdf', 'doc', 'docx', 'xls', 'xlsx','csv'])
    target = ['.' + j for j in target0]
    mingxi = []
    search_file(fname, start_dir, target, keyname)
    with open(fname, 'r', encoding='utf-8', errors='replace')  as f:
        for j in f.readlines():
            mingxi.append(j)

    fmingxi = open(fname, 'w', encoding='utf-8')
    fmingxi.writelines(mingxi)
    fmingxi.close()
else :
    msg = '请点选文件'
    file = easygui.fileopenbox(msg)
    fname,file = searchSingleFile(fname,file)
os.startfile(fname)


'''
import codecs
with codecs.open(r"F:\a00nutstore\my note\typoraStudy\你好 Typora.md", "rb", 'utf-8', errors='ignore') as txtfile:
    for line in txtfile:
        print(line)
'''






